"""report_generator.py — Generate a professional Word document variance report.

Public API:
    from report_generator import generate_report

    generate_report(
        full_report=df,          # from variance_calculator (all accounts)
        commentary=comments,      # from commentary_generator (may be [])
        config=cfg,              # client config dict (from load_config)
        output_path="output/report.docx",
        period_label="April 2026",
    )

CLI (standalone sanity check):
    python report_generator.py --output output/sample.docx --period "April 2026"
"""

import argparse
import logging
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_FONT = "Calibri"

_C = {
    "navy":       "1F497D",
    "white":      "FFFFFF",
    "red_fill":   "FFC7CE",
    "red_text":   "9C0006",
    "yel_fill":   "FFEB9C",
    "yel_text":   "9C5700",
    "logo_bg":    "E0E0E0",
    "logo_text":  "808080",
    "kpi_bg":     "EBF3FB",
    "alt_row":    "F2F2F2",
    "comm_bg":    "F3F3F3",
    "grey_text":  "666666",
    "light_grey": "CCCCCC",
    "green_text": "2E7D32",
}

_MARGIN = Inches(0.75)

_VAR_HEADERS = [
    "Account #", "Account Name", "Current ($)", "Prior ($)",
    "Variance ($)", "Variance %", "Flag",
]
_VAR_WIDTHS = [
    Inches(0.65), Inches(1.85), Inches(1.00), Inches(1.00),
    Inches(1.00), Inches(0.75), Inches(0.55),
]

_APP_HEADERS = [
    "Account #", "Account Name", "Current ($)", "Prior ($)",
    "Variance ($)", "Var %",
]
_APP_WIDTHS = [
    Inches(0.70), Inches(2.10), Inches(1.10), Inches(1.10),
    Inches(1.10), Inches(0.75),
]


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_report(
    full_report: pd.DataFrame,
    commentary: list,
    config: dict,
    output_path: str,
    period_label: str = "",
) -> None:
    """Generate a professional Word document variance report.

    Sections produced (in order):
        1. Header block  — logo placeholder, client name, period, preparer.
        2. Executive Summary — KPI cards: total accounts, flagged count,
           largest variance.
        3. Variance Detail Table — all flagged accounts, colour-coded by
           variance percentage (red > 10 %, yellow 5–10 %).
        4. AI Commentary — per-account explanations keyed from *commentary*.
           Omitted entirely when *commentary* is empty.
        5. Appendix — full trial balance with alternating row fills.

    Running header (page 2+): "Firm | Client Name".
    Footer: "Page X of Y" right-aligned.

    Args:
        full_report:  DataFrame produced by variance_calculator with columns
                      Account_Number, Account_Name, Current_Balance,
                      Prior_Balance, Variance_Amount, Variance_Pct, Flagged.
        commentary:   List of dicts from commentary_generator, each with
                      account_number, account_name, commentary. Pass [] to
                      skip the AI Commentary section.
        config:       Client configuration dict. Keys used: client_name,
                      preparer, firm, variance_threshold_pct,
                      variance_threshold_abs.
        output_path:  Destination .docx path. Parent directories are created.
        period_label: Human-readable period string, e.g. "April 2026".

    Raises:
        ValueError: If full_report is missing required columns.
        OSError:    If the output file cannot be written (e.g. open in Word).
    """
    _validate_columns(full_report)

    flagged = full_report[full_report["Flagged"]].copy().reset_index(drop=True)
    commentary_map = {c["account_number"]: c["commentary"] for c in commentary}

    doc = Document()
    _configure_doc(doc, config)

    _write_header_section(doc, config, period_label)
    _write_executive_summary(doc, full_report, flagged, config)

    doc.add_page_break()
    _write_variance_table(doc, flagged)

    if commentary:
        doc.add_page_break()
        _write_commentary_section(doc, flagged, commentary_map)

    doc.add_page_break()
    _write_appendix(doc, full_report)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    log.info("Report saved: '%s' (%d accounts, %d flagged).", path, len(full_report), len(flagged))


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------

def _configure_doc(doc: Document, config: dict) -> None:
    """Set page margins, default font, and running header/footer on every section."""
    for section in doc.sections:
        section.top_margin    = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin   = _MARGIN
        section.right_margin  = _MARGIN
        section.different_first_page_header_footer = True  # suppress header on p.1
        _setup_running_header(section, config)
        _setup_footer(section)

    doc.styles["Normal"].font.name = _FONT
    doc.styles["Normal"].font.size = Pt(10)


def _setup_running_header(section, config: dict) -> None:
    """Add "Firm | Client" text with a bottom rule to pages 2+."""
    firm   = config.get("firm", "")
    client = config.get("client_name", "")

    p = section.header.paragraphs[0]
    p.clear()
    run = p.add_run(f"{firm}  |  {client}")
    run.font.name  = _FONT
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(*_hex("grey_text"))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_bottom_border(p, _C["light_grey"])


def _setup_footer(section) -> None:
    """Add right-aligned 'Page X of Y' footer."""
    p = section.footer.paragraphs[0]
    p.clear()
    _text_run(p, "Page ", 8, "grey_text")
    _field_run(p, "PAGE", 8, "grey_text")
    _text_run(p, " of ", 8, "grey_text")
    _field_run(p, "NUMPAGES", 8, "grey_text")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ---------------------------------------------------------------------------
# Section writers
# ---------------------------------------------------------------------------

def _write_header_section(doc: Document, config: dict, period_label: str) -> None:
    """Page-1 header: logo placeholder (left) + report metadata (right)."""
    table = doc.add_table(rows=1, cols=2)
    _kill_borders(table)
    row = table.rows[0]

    # ── Left cell: logo placeholder ──────────────────────────────────────────
    logo = row.cells[0]
    logo.width = Inches(1.5)
    _shade(logo, "logo_bg")
    _valign(logo, "center")

    logo.paragraphs[0].clear()
    _blank_para(logo)
    p = logo.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ COMPANY\nLOGO ]")
    r.font.name  = _FONT
    r.font.size  = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(*_hex("logo_text"))
    _blank_para(logo)

    # ── Right cell: report metadata ──────────────────────────────────────────
    info = row.cells[1]
    _shade(info, "navy")
    _valign(info, "center")

    lines = [
        ("Month-End Close Report",                                              18, True),
        (config.get("client_name", ""),                                         13, False),
        (f"Period: {period_label}",                                             11, False),
        (f"Prepared by: {config.get('preparer','')}  |  {config.get('firm','')}", 9, False),
        (f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",              8, False),
    ]
    for i, (text, size, bold) in enumerate(lines):
        p = info.paragraphs[0] if i == 0 else info.add_paragraph()
        p.clear()
        r = p.add_run(text)
        r.font.name  = _FONT
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(8 if i == 0 else 3)
        p.paragraph_format.space_after  = Pt(8 if i == len(lines) - 1 else 2)

    doc.add_paragraph()


def _write_executive_summary(
    doc: Document,
    full_report: pd.DataFrame,
    flagged: pd.DataFrame,
    config: dict,
) -> None:
    """Three KPI tiles: total accounts, flagged count, largest variance."""
    _section_heading(doc, "Executive Summary")

    total        = len(full_report)
    flagged_n    = len(flagged)
    largest_abs  = flagged["Variance_Amount"].abs().max() if not flagged.empty else 0.0
    largest_name = (
        flagged.loc[flagged["Variance_Amount"].abs().idxmax(), "Account_Name"]
        if not flagged.empty else "—"
    )
    thr_pct = config.get("variance_threshold_pct", 5)
    thr_abs = config.get("variance_threshold_abs", 5000)

    kpis = [
        ("Total Accounts Reviewed", str(total),          "All accounts processed"),
        ("Flagged Variances",        str(flagged_n),      f"Threshold: ≥{thr_pct}%  |  ≥{_curr(thr_abs)}"),
        ("Largest Variance (Abs.)",  _curr(largest_abs),  largest_name),
    ]

    table = doc.add_table(rows=2, cols=3)
    _kill_borders(table)
    _col_widths(table, [Inches(2.2), Inches(2.2), Inches(2.4)])

    for i, (label, value, sub) in enumerate(kpis):
        # Label row (navy banner)
        lc = table.rows[0].cells[i]
        _shade(lc, "navy")
        p = lc.paragraphs[0]
        p.clear()
        r = p.add_run(label)
        r.font.name = _FONT; r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)

        # Value row (light blue background)
        vc = table.rows[1].cells[i]
        _shade(vc, "kpi_bg")
        p2 = vc.paragraphs[0]
        p2.clear()
        r2 = p2.add_run(value)
        r2.font.name = _FONT; r2.font.size = Pt(18); r2.font.bold = True
        r2.font.color.rgb = RGBColor(*_hex("navy"))
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after  = Pt(2)

        p3 = vc.add_paragraph()
        r3 = p3.add_run(sub)
        r3.font.name = _FONT; r3.font.size = Pt(8); r3.font.italic = True
        r3.font.color.rgb = RGBColor(*_hex("grey_text"))
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(8)

    doc.add_paragraph()


def _write_variance_table(doc: Document, flagged: pd.DataFrame) -> None:
    """Flagged-accounts detail table with red/yellow colour-coded rows."""
    _section_heading(doc, "Variance Detail — Flagged Accounts")

    if flagged.empty:
        p = doc.add_paragraph("No accounts exceeded the variance thresholds for this period.")
        p.runs[0].font.italic = True
        return

    table = doc.add_table(rows=1 + len(flagged), cols=len(_VAR_HEADERS))
    table.style = "Table Grid"
    _col_widths(table, _VAR_WIDTHS)

    # Header row
    for i, h in enumerate(_VAR_HEADERS):
        _cell(table.rows[0].cells[i], h,
              size=9, bold=True,
              fg="white", bg="navy",
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, (_, row) in enumerate(flagged.iterrows(), start=1):
        pct = abs(float(row["Variance_Pct"]))
        if pct > 10:
            bg, fg = "red_fill", "red_text"
        elif pct > 5:
            bg, fg = "yel_fill", "yel_text"
        else:
            bg, fg = None, None

        cells = table.rows[ri].cells
        data = [
            (str(row["Account_Number"]),        WD_ALIGN_PARAGRAPH.CENTER),
            (str(row["Account_Name"]),           WD_ALIGN_PARAGRAPH.LEFT),
            (_curr(row["Current_Balance"]),      WD_ALIGN_PARAGRAPH.RIGHT),
            (_curr(row["Prior_Balance"]),         WD_ALIGN_PARAGRAPH.RIGHT),
            (_curr(row["Variance_Amount"]),       WD_ALIGN_PARAGRAPH.RIGHT),
            (_pct(row["Variance_Pct"]),           WD_ALIGN_PARAGRAPH.RIGHT),
            ("FLAG" if row["Flagged"] else "",    WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for ci, (text, align) in enumerate(data):
            _cell(cells[ci], text,
                  size=9, bold=(bg is not None),
                  fg=fg, bg=bg, align=align)

    doc.add_paragraph()


def _write_commentary_section(
    doc: Document,
    flagged: pd.DataFrame,
    commentary_map: dict,
) -> None:
    """AI-generated 2-3 sentence explanations, one per flagged account."""
    _section_heading(doc, "AI-Generated Commentary")

    intro = doc.add_paragraph()
    r = intro.add_run(
        "The following explanations were generated by an AI model to highlight common "
        "business drivers for each flagged variance. Review and validate before including "
        "in client deliverables."
    )
    r.font.name = _FONT; r.font.size = Pt(9); r.font.italic = True
    r.font.color.rgb = RGBColor(*_hex("grey_text"))
    intro.paragraph_format.space_after = Pt(10)

    for _, row in flagged.iterrows():
        acct_num  = str(row["Account_Number"])
        acct_name = str(row["Account_Name"])
        var_amt   = float(row["Variance_Amount"])
        var_pct   = float(row["Variance_Pct"])
        text      = commentary_map.get(
            acct_num,
            "[No AI commentary available for this account.]",
        )

        # Account sub-heading
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(10)
        ph.paragraph_format.space_after  = Pt(2)
        rh = ph.add_run(f"{acct_name}  ({acct_num})")
        rh.font.name = _FONT; rh.font.size = Pt(10); rh.font.bold = True
        rh.font.color.rgb = RGBColor(*_hex("navy"))

        # Variance summary line
        direction = "▲" if var_amt >= 0 else "▼"
        color_key = "green_text" if var_amt >= 0 else "red_text"
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after  = Pt(4)
        rs = ps.add_run(f"{direction}  {_curr(abs(var_amt))}  ({_pct(var_pct)})")
        rs.font.name = _FONT; rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(*_hex(color_key))

        # Commentary in a shaded call-out box (borderless 1-cell table)
        tbl = doc.add_table(rows=1, cols=1)
        _kill_borders(tbl)
        box = tbl.rows[0].cells[0]
        _shade(box, "comm_bg")
        pb = box.paragraphs[0]
        pb.clear()
        rb = pb.add_run(text)
        rb.font.name = _FONT; rb.font.size = Pt(9)
        pb.paragraph_format.left_indent   = Inches(0.12)
        pb.paragraph_format.right_indent  = Inches(0.12)
        pb.paragraph_format.space_before  = Pt(5)
        pb.paragraph_format.space_after   = Pt(5)

        doc.add_paragraph()


def _write_appendix(doc: Document, full_report: pd.DataFrame) -> None:
    """Full trial balance — all accounts in a compact, zebra-striped table."""
    _section_heading(doc, "Appendix — Full Trial Balance")

    note = doc.add_paragraph()
    rn = note.add_run(
        f"All {len(full_report)} accounts for this period. "
        "Flagged accounts are highlighted."
    )
    rn.font.name = _FONT; rn.font.size = Pt(9); rn.font.italic = True
    rn.font.color.rgb = RGBColor(*_hex("grey_text"))
    note.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1 + len(full_report), cols=len(_APP_HEADERS))
    table.style = "Table Grid"
    _col_widths(table, _APP_WIDTHS)

    # Header
    for i, h in enumerate(_APP_HEADERS):
        _cell(table.rows[0].cells[i], h,
              size=8, bold=True,
              fg="white", bg="navy",
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, (_, row) in enumerate(full_report.iterrows(), start=1):
        is_flagged = bool(row["Flagged"])
        pct        = abs(float(row["Variance_Pct"]))

        if is_flagged and pct > 10:
            bg, fg = "red_fill", "red_text"
        elif is_flagged and pct > 5:
            bg, fg = "yel_fill", "yel_text"
        elif ri % 2 == 0:
            bg, fg = "alt_row", None
        else:
            bg, fg = None, None

        data = [
            (str(row["Account_Number"]),       WD_ALIGN_PARAGRAPH.CENTER),
            (str(row["Account_Name"]),          WD_ALIGN_PARAGRAPH.LEFT),
            (_curr(row["Current_Balance"]),     WD_ALIGN_PARAGRAPH.RIGHT),
            (_curr(row["Prior_Balance"]),        WD_ALIGN_PARAGRAPH.RIGHT),
            (_curr(row["Variance_Amount"]),      WD_ALIGN_PARAGRAPH.RIGHT),
            (_pct(row["Variance_Pct"]),          WD_ALIGN_PARAGRAPH.RIGHT),
        ]
        for ci, (text, align) in enumerate(data):
            _cell(table.rows[ri].cells[ci], text,
                  size=8, fg=fg, bg=bg, align=align)


# ---------------------------------------------------------------------------
# Low-level XML / styling helpers
# ---------------------------------------------------------------------------

def _validate_columns(df: pd.DataFrame) -> None:
    """Raise ValueError listing every missing required column."""
    required = {
        "Account_Number", "Account_Name", "Current_Balance",
        "Prior_Balance", "Variance_Amount", "Variance_Pct", "Flagged",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(
            f"full_report is missing required column(s): {sorted(missing)}. "
            f"Found: {list(df.columns)}"
        )


def _section_heading(doc: Document, text: str) -> None:
    """Bold navy heading paragraph with a thin navy underline rule."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = _FONT; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(*_hex("navy"))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    _para_bottom_border(p, _C["navy"])


def _cell(
    cell,
    text: str,
    size: int = 10,
    bold: bool = False,
    italic: bool = False,
    fg: str = None,
    bg: str = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Write text into a table cell with optional fill and font styling."""
    if bg:
        _shade(cell, bg)
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(str(text))
    r.font.name   = _FONT
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if fg:
        r.font.color.rgb = RGBColor(*_hex(fg))


def _shade(cell, color_key: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc  = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _C[color_key])
    tcPr.append(shd)


def _valign(cell, alignment: str = "center") -> None:
    """Set vertical alignment on a table cell (top, center, bottom)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    va   = OxmlElement("w:vAlign")
    va.set(qn("w:val"), alignment)
    tcPr.append(va)


def _kill_borders(table) -> None:
    """Remove all borders from a table (overrides any style-inherited borders)."""
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    table._tbl.tblPr.append(tblBorders)


def _col_widths(table, widths: list) -> None:
    """Set per-column widths (EMU) across all rows for consistent layout."""
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w


def _para_bottom_border(paragraph, color: str = "CCCCCC") -> None:
    """Draw a thin bottom rule under a paragraph."""
    pPr    = paragraph._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _blank_para(cell) -> None:
    """Add an empty padding paragraph inside a cell."""
    p = cell.add_paragraph()
    r = p.add_run()
    r.font.size = Pt(6)


def _text_run(paragraph, text: str, size: int, color_key: str = None) -> None:
    """Append a styled text run to an existing paragraph."""
    r = paragraph.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size)
    if color_key:
        r.font.color.rgb = RGBColor(*_hex(color_key))


def _field_run(paragraph, field_code: str, size: int, color_key: str = None) -> None:
    """Append a Word field (PAGE, NUMPAGES, etc.) to an existing paragraph."""
    run = paragraph.add_run()
    run.font.name = _FONT
    run.font.size = Pt(size)
    if color_key:
        run.font.color.rgb = RGBColor(*_hex(color_key))

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _hex(color_key: str) -> tuple:
    """Return (R, G, B) bytes tuple for a palette key."""
    return tuple(bytes.fromhex(_C[color_key]))


def _curr(val) -> str:
    """Format a number as accounting-style currency: $(1,234.56) for negatives."""
    v = float(val)
    return f"(${abs(v):,.2f})" if v < 0 else f"${v:,.2f}"


def _pct(val) -> str:
    """Format a number as a percentage string with one decimal place."""
    return f"{float(val):.1f}%"


# ---------------------------------------------------------------------------
# Standalone execution — generates a sample .docx without real API calls
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    logging.basicConfig(level=logging.DEBUG, format="%(levelname)-8s %(name)s: %(message)s")

    parser = argparse.ArgumentParser(description="Generate a sample Word variance report.")
    parser.add_argument("--output", default="output/sample_report.docx")
    parser.add_argument("--period", default="April 2026")
    args = parser.parse_args()

    sample_report = pd.DataFrame({
        "Account_Number":  ["4000", "4100", "5000", "6000", "6100", "7000"],
        "Account_Name":    ["Product Sales", "Service Revenue", "Materials",
                            "Salaries", "Rent", "Interest Income"],
        "Current_Balance": [178_000, 46_500, 72_000, 42_000, 8_000,  1_200],
        "Prior_Balance":   [150_000, 45_000, 60_000, 35_000, 8_000,  1_000],
        "Variance_Amount": [ 28_000,  1_500, 12_000,  7_000,     0,    200],
        "Variance_Pct":    [  18.67,   3.33,  20.00,  20.00,  0.00,  20.00],
        "Flagged":         [   True,  False,   True,   True,  False,   True],
    })

    sample_commentary = [
        {
            "account_number": "4000",
            "account_name": "Product Sales",
            "commentary": (
                "Product Sales increased $28,000 (18.7%) driven by strong seasonal demand "
                "and completion of a deferred Q1 contract. The uptick aligns with "
                "historical Q2 patterns observed in prior years."
            ),
        },
        {
            "account_number": "5000",
            "account_name": "Materials",
            "commentary": (
                "Materials expense rose $12,000 (20.0%) consistent with higher sales volume "
                "and an 8% supplier price increase effective April 1. "
                "The variance is proportional to production activity and within budget."
            ),
        },
        {
            "account_number": "6000",
            "account_name": "Salaries",
            "commentary": (
                "Salaries increased $7,000 (20.0%) due to annual merit adjustments "
                "processed in April and the addition of one part-time contractor. "
                "This variance was anticipated per the approved headcount budget."
            ),
        },
        {
            "account_number": "7000",
            "account_name": "Interest Income",
            "commentary": (
                "Interest Income increased $200 (20.0%) reflecting higher money-market "
                "yields following the Fed rate adjustment in Q1. "
                "This positive variance is expected to continue through Q3."
            ),
        },
    ]

    sample_config = {
        "client_name": "Acme Corporation",
        "preparer": "Jane Smith, CPA",
        "firm": "Smith & Associates",
        "variance_threshold_pct": 10,
        "variance_threshold_abs": 5000,
    }

    generate_report(
        full_report=sample_report,
        commentary=sample_commentary,
        config=sample_config,
        output_path=args.output,
        period_label=args.period,
    )
    print(f"Report generated: {args.output}")
